from typing import List
from langchain_core.documents import Document
import pdfplumber
import docx
import os

def load_txt_file(file_path: str) -> List[Document]:
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()
    return [Document(page_content=text, metadata={"source": os.path.basename(file_path)})]

def load_pdf_file(file_path: str) -> List[Document]:
    docs = []
    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text:
                docs.append(Document(
                    page_content=text,
                    metadata={"source": os.path.basename(file_path), "page": i + 1}
                ))
    return docs

def load_docx_file(file_path: str) -> List[Document]:
    doc = docx.Document(file_path)
    full_text = "\n".join([para.text for para in doc.paragraphs])
    return [Document(page_content=full_text, metadata={"source": os.path.basename(file_path)})]

def parse_file(file_path: str) -> List[Document]:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".txt":
        return load_txt_file(file_path)
    elif ext == ".pdf":
        return load_pdf_file(file_path)
    elif ext == ".docx":
        return load_docx_file(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
